# create a reverse index from a docx document of paragraphs
# using a specific convention to select paragraphs

from docx import Document
from itertools import ifilter
from itertools import groupby
import re
import utils
import logging
import raw_indexer
import elaborated_indexer
from index_element import IndexElement

log = logging.getLogger()

# Constants
INDEX_LENGTH_THRESHOLD = 50
REGEX_UNWANTED_CHAR = re.compile(ur"[\u00ab\u00bb]+")
REGEX_QUOTATION_MARKS = re.compile(ur"[\u2019\u0027\u201c\u201d\u05f3]+")
REGEX_MULTIPLE_SPACES = re.compile(" +")


# Apply filter on paragraphs to select only the area to index
def filter_paragraph(paragraphs,paragraph_style):

    def filter(p):
        return p.style.name == paragraph_style and p.text.strip() != ""
    
    return ifilter(filter,paragraphs)

# Reduction function that marks true if the current paragraph should be considered for numeritation
def numeritation_function(paragraph):
    if paragraph._element.pPr.numPr is None: # Try to detect unnumbered paragraphs
        return False
    return True

# From a list of paragraph, return a list of text corresponding to the numeritation function changes
def compress_and_text(pars):
    compressed = list()
    for p in pars:
        if numeritation_function(p):
            compressed.append(p.text)
        elif len(compressed) > 0:
            compressed[-1] = compressed[-1] + "\n" + p.text
        else:
            log.info("Skipping text as before first numeritation: %s" % p.text[:20]+"..." )
    return compressed

# Produce a debugging file to understand where the numeration have gaps 
def debug_numeration(docx_filename, paragraph_style):

    from docx.enum.text import WD_COLOR_INDEX

    def reset_highlight(doc):
        for p in doc.paragraphs:
            for r in p.runs:
                r.font.highlight_color = None

    def invert_highlight(doc):
        for p in doc.paragraphs:
            for r in p.runs:
                if r.font.highlight_color is None:
                    r.font.highlight_color =  WD_COLOR_INDEX.YELLOW
                else:
                     r.font.highlight_color = None     
    
    log.info("Elaborating from "+docx_filename)
    doc = Document(docx_filename)

    log.info("Resetting Highlight")
    reset_highlight(doc)

    log.info("Adding Numeration and highlight")
    # selecing paragraph that are candidate for numeration
    # Add numerotation in the way is calculated after and format highlight
    num_index = 1
    for p in filter_paragraph(doc.paragraphs,paragraph_style):
        if numeritation_function(p):
                p.text = ("(%s)" % num_index ) + p.text
                num_index=num_index+1
        # add highlight to be invered later
        for r in p.runs:
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    invert_highlight(doc)
    log.info("Saving result as .debug.docx")
    doc.save(docx_filename+".debug.docx")



# Filter paragraph and import 
def import_paragraphs(docx_filename, paragraph_style):

    def clean_text(text):
        text = REGEX_UNWANTED_CHAR.sub("",text)    #remove bad chars
        text = REGEX_QUOTATION_MARKS.sub("'",text) #replace bad apostrophe with more compliant to facilitate parsing
        text = REGEX_MULTIPLE_SPACES.sub(" ",text)
        return text.lower()

    log.info("Loading from %s",docx_filename)
    doc = Document(docx_filename)
    listpars = compress_and_text(filter_paragraph(doc.paragraphs,paragraph_style))
    cleaned_listpars = map(clean_text,listpars)
    return cleaned_listpars


# Consolidate together all the index elements that have the same entry
def consolidate_index(raw_reverse_index):
    log.info("Consolidating Index")
    consolidated_index = list()
    sorted_raw_index = sorted(raw_reverse_index, key=lambda x:x.entry)  # sort to have the groupBy semantics working
    for entry, group in groupby(sorted_raw_index, lambda x: x.entry):
        consolidated_entry = reduce((lambda x, y: x + y), group)
        consolidated_index.append(consolidated_entry)
    return consolidated_index


# create the reverse index and output it as a csv file
def create_raw_index(docx_filename,output_filename,paragraph_style):
    doc = import_paragraphs(docx_filename,paragraph_style)
    base_index = raw_indexer.parse_and_split(doc)
    index = consolidate_index(base_index)
    utils.write_csv(output_filename,index)
    return 

# create the reverse index and output it as a csv file
def create_structured_index(docx_filename,output_filename,paragraph_style,structure_file):
    doc = import_paragraphs(docx_filename,paragraph_style)
    index_structure = elaborated_indexer.parse_yaml(structure_file)
    base_index = elaborated_indexer.parse_and_split(index_structure,doc)
    index = consolidate_index(base_index)
    utils.write_string(output_filename,elaborated_indexer.index_printer(index))
    return 


